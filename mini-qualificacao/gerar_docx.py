"""Gera a mini-qualificacao (Arial 12, max. 5 paginas) a partir do template oficial
e dos insumos numericos produzidos pelos notebooks 01-05.

Uso:  python mini-qualificacao/gerar_docx.py
Saida: mini-qualificacao/INF-009 2026.2 - Mini-qualificacao - Priscila Santos.docx
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
MQ = ROOT / "mini-qualificacao"
INS = MQ / "insumos_gerados"
TEMPLATE = MQ / "INF-009 2026.2 - Template Mini-qualificacao.docx"
SAIDA = MQ / "INF-009 2026.2 - Mini-qualificacao - Priscila Santos.docx"

FONTE = "Arial"
TAM = Pt(12)
TAM_TAB = Pt(12)  # documento inteiro em Arial 12, inclusive tabela e referencias


# --------------------------------------------------------------------------
# Infraestrutura de formatacao
# --------------------------------------------------------------------------
def _set_run(run, size=TAM, bold=False, italic=False):
    run.font.name = FONTE
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONTE)


def limpar_corpo(doc):
    """Remove todo o conteudo do corpo, preservando cabecalho/rodape e sectPr."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def doc_defaults(doc):
    """Arial 12 tambem no docDefaults, para que ate as marcas estruturais de
    linha de tabela (paragrafos vazios de fim de linha) herdem o tamanho certo."""
    styles = doc.styles.element
    dd = styles.find(qn("w:docDefaults"))
    if dd is None:
        return
    rpr_dflt = dd.find(qn("w:rPrDefault"))
    if rpr_dflt is None:
        return
    rpr = rpr_dflt.find(qn("w:rPr"))
    if rpr is None:
        rpr = rpr_dflt.makeelement(qn("w:rPr"), {})
        rpr_dflt.append(rpr)
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONTE)
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = rpr.makeelement(qn(tag), {})
            rpr.append(el)
        el.set(qn("w:val"), str(int(TAM.pt * 2)))  # meio-pontos


def normal_style(doc):
    doc_defaults(doc)
    st = doc.styles["Normal"]
    st.font.name = FONTE
    st.font.size = TAM
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), FONTE)
    pf = st.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0


def par(doc, texto="", bold=False, italic=False, align="just", size=TAM,
        space_after=4, space_before=0, esquerda=0.0, primeira_linha=0.0):
    p = doc.add_paragraph()
    p.alignment = {
        "just": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }[align]
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.0
    if esquerda:
        pf.left_indent = Cm(esquerda)
    if primeira_linha:
        pf.first_line_indent = Cm(primeira_linha)
    if texto:
        _set_run(p.add_run(texto), size=size, bold=bold, italic=italic)
    return p


def rico(doc, partes, align="just", space_after=4, esquerda=0.0, size=TAM):
    """partes = lista de (texto, bold) para negrito seletivo dentro do paragrafo."""
    p = doc.add_paragraph()
    p.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if esquerda:
        pf.left_indent = Cm(esquerda)
    for texto, bold in partes:
        _set_run(p.add_run(texto), size=size, bold=bold)
    return p


def titulo_secao(doc, texto):
    return par(doc, texto, bold=True, align="left", space_before=8, space_after=4)


def _bordas(t):
    """O template nao traz o estilo 'Table Grid'; aplica bordas via XML."""
    tbl_pr = t._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{lado}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tbl_pr.append(borders)


def tabela(doc, cabecalho, linhas, larguras=None, size=TAM_TAB):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    _bordas(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(cabecalho):
        p = hdr[i].paragraphs[0]
        for r in list(p.runs):  # remove o run vazio que herdaria o tamanho do template
            r._element.getparent().remove(r._element)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        _set_run(p.add_run(h), size=size, bold=True)
    for linha in linhas:
        cells = t.add_row().cells
        for i, v in enumerate(linha):
            p = cells[i].paragraphs[0]
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            _set_run(p.add_run(str(v)), size=size)
    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Cm(w)
    return t


# --------------------------------------------------------------------------
# Carga dos insumos gerados pelos notebooks
# --------------------------------------------------------------------------
def carregar_insumos():
    d = {}
    d["consolidado"] = pd.read_csv(INS / "tabela_resultados_consolidados.csv")
    d["boot"] = pd.read_csv(INS / "tabela_bootstrap_ic95.csv")
    d["binario"] = pd.read_csv(INS / "tabela_objetivo2_binario.csv")
    d["decisao"] = pd.read_csv(INS / "decisao_hipotese.csv")
    d["importancia"] = pd.read_csv(INS / "tabela_importancia_features.csv", index_col=0)
    d["estabilidade"] = pd.read_csv(INS / "tabela_estabilidade_features.csv", index_col=0)
    return d


def n(x, casas=3):
    """Formata número no padrão brasileiro (vírgula decimal)."""
    return f"{x:.{casas}f}".replace(".", ",")


def pct(x, casas=2):
    return f"{100*x:.{casas}f}".replace(".", ",") + "%"


def mil(x):
    """Separador de milhar no padrão brasileiro: 7918 -> 7.918."""
    return f"{int(x):,}".replace(",", ".")


REFERENCIAS = [
    "GUPTA, A. et al. DAiSEE: Towards User Engagement Recognition in the Wild. arXiv:1609.01885, 2016.",
    "LEE, H. et al. Investigating the Effects of a Real-time Student Monitoring Interface on Instructors' "
    "Monitoring Practices in Online Teaching. 2024.",
    "NASCIMENTO, M. D. et al. Computational Vision for Analyzing Student Emotions in E-learning: a Preliminary "
    "Study. 2023.",
    "MRAYHI, S. et al. Enhancing MOOCs through Real-time Learner Engagement and Emotion Detection Using Computer "
    "Vision and Machine Learning. 2024.",
    "ABEDI, A.; KHAN, S. S. Affect-driven Ordinal Engagement Measurement from Video. arXiv:2106.10882, 2021.",
    "ABEDI, A.; KHAN, S. S. Engagement Measurement Based on Facial Landmarks and Spatial-Temporal Graph "
    "Convolutional Networks. arXiv:2403.17175, 2024.",
    "LIAO, J.; LIANG, Y.; PAN, J. Deep Facial Spatiotemporal Network for Engagement Prediction in Online Learning. "
    "Applied Intelligence, 2021. DOI:10.1007/s10489-020-02139-8.",
    "MEHTA, N. K. et al. Three-dimensional DenseNet Self-attention Neural Network for Automatic Detection of "
    "Student's Engagement. Applied Intelligence, 2022. DOI:10.1007/s10489-022-03200-4.",
    "ABEDI, A. et al. Bag of States: A Non-sequential Approach to Video-based Engagement Measurement. "
    "arXiv:2301.06730, 2023.",
    "SUGIHDHARMA, M. R.; BACHTIAR, F. A. Engagement and Affective State Detection Using Scaled Dot-Product "
    "Attention CNN on Eye Landmarks and Gaze. 2023. DOI:10.1145/3626641.3626938.",
    "SINGH, A. et al. VisioPhysioENet: Multimodal Engagement Detection Using Visual and Physiological Signals. "
    "arXiv:2409.16126, 2024.",
    "MANDIA, S. et al. EngageFormer: Multi-view Transformer for Student Engagement Recognition. "
    "arXiv:2502.10813, 2025.",
    "GOTHWAL, A. et al. ViBED-Net: Video-Based Engagement Detection Network. arXiv:2510.18016, 2025.",
    "SAFA, M.; ABEDI, A.; KHAN, S. S. Supervised Contrastive Ordinal Learning for Engagement Measurement. "
    "arXiv:2505.20676, 2025.",
    "GOYAL, N. et al. Zero-Shot Vision-Language Models for Student Engagement Recognition: A Benchmark. "
    "arXiv:2606.21861, 2026.",
    "CAO, W.; MIRJALILI, V.; RASCHKA, S. Rank Consistent Ordinal Regression for Neural Networks (CORAL). "
    "Pattern Recognition Letters, v. 140, p. 325-331, 2020.",
    "LUGARESI, C. et al. MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172, 2019.",
]


def montar_documento(d):
    doc = Document(str(TEMPLATE))
    limpar_corpo(doc)
    normal_style(doc)

    dec = d["decisao"].iloc[0]
    cons = d["consolidado"]
    boot = d["boot"]
    binr = d["binario"].iloc[0]
    imp = d["importancia"]
    est = d["estabilidade"]
    meta = d["meta"]

    base_row = cons[cons["modelo"].str.contains("trivial", case=False)].iloc[0]
    melhor_row = cons.iloc[0]
    b_model = boot.iloc[0]
    b_base = boot.iloc[1]

    veredito = dec["veredito_hipotese"]
    top5 = list(imp.index[:5])
    estaveis = est[est.iloc[:, 0] >= 4].index.tolist()

    # ---------------- Folha de rosto ----------------
    par(doc, "Universidade Federal do ABC", align="center", space_after=0)
    par(doc, "Pós-graduação em Engenharia da Informação", align="center", space_after=0)
    par(doc, "INF-009 – Projeto e Comunicação de Pesquisa em Engenharia da Informação",
        align="center", space_after=0)
    par(doc, "Quadrimestre: 2026.2", align="center", space_after=10)
    par(doc, "Classificação de Engajamento em EaD com Visão Computacional",
        bold=True, align="center", space_after=10)
    par(doc, "Priscila Henrique Medeiro dos Santos (priscilahms@yahoo.com.br)",
        align="center", space_after=0)
    par(doc, "Santo André, agosto de 2026", align="center", space_after=8)

    # ---------------- 1. Resumo ----------------
    titulo_secao(doc, "1. Resumo")
    par(doc,
        "O acompanhamento do engajamento discente em aulas síncronas de Educação a Distância (EaD) esbarra nos "
        "limites da percepção humana: o docente não observa simultaneamente dezenas de participantes, o que "
        "retarda a detecção de desengajamento e alimenta a evasão [2][3][4]. No dataset de referência da área, o "
        "DAiSEE [1], a literatura reporta majoritariamente acurácia como métrica única, subtratando um "
        "desbalanceamento extremo (as duas classes de baixo engajamento somam menos de 6% dos clipes) e a natureza "
        "ordinal do rótulo. Este trabalho entrega uma Prova de Conceito leve e preservadora de privacidade: marcos "
        "faciais, íris/gaze e pose da cabeça extraídos com MediaPipe FaceMesh [17], agregados por clipe e "
        "classificados por modelos clássicos (Trilha A) e por uma rede temporal rasa BiLSTM (Trilha B), com "
        "comparação explícita de estratégias de balanceamento e da formulação ordinal CORAL [16]. A avaliação usa "
        "os splits oficiais disjuntos por sujeito e um protocolo multimétrica (acurácia, macro-F1, kappa de Cohen, "
        "revocação por classe) com intervalos de confiança por bootstrap, confrontando o modelo com o baseline "
        f"trivial de classe majoritária. O melhor modelo atingiu macro-F1 de {n(melhor_row['macro_f1'])} e kappa de "
        f"{n(melhor_row['kappa'])} contra {n(base_row['macro_f1'])} e {n(base_row['kappa'])} do baseline, "
        f"levando a hipótese a ser considerada {veredito.lower()}.")

    # ---------------- 2. Introducao ----------------
    titulo_secao(doc, "2. Introdução e Motivação")
    par(doc,
        "Em ambientes virtuais de aprendizagem, o professor perde os sinais não verbais que, na sala presencial, "
        "sinalizam tédio, confusão ou dispersão. A literatura associa essa cegueira situacional à queda de "
        "motivação e a elevadas taxas de evasão [3][4]. Automatizar a leitura desses sinais a partir do vídeo já "
        "capturado pelas webcams é, portanto, uma intervenção de baixo custo marginal com potencial pedagógico "
        "direto. O desafio é que engajamento não é uma emoção básica: é um estado acadêmico difuso, anotado de "
        "forma ordinal e altamente desbalanceada, e observado “in the wild”, com iluminação, enquadramento e "
        "qualidade heterogêneos.")
    rico(doc, [
        ("Questões de pesquisa. ", True),
        ("(QP1) Quais indicadores visuais extraíveis em tempo real – abertura ocular, direção do olhar, pose da "
         "cabeça, movimentação da boca – estão mais associados aos níveis de engajamento? (QP2) Modelos treinados "
         "apenas sobre marcos faciais, sem transmitir ou armazenar a imagem do rosto, conseguem separar estudantes "
         "engajados de desengajados em vídeo real com desempenho superior a um classificador trivial?", False),
    ])
    rico(doc, [
        ("Hipótese. ", True),
        ("Modelos de visão computacional que analisam continuamente indicadores comportamentais (expressões "
         "faciais, orientação da cabeça, comportamentos de distração) permitem identificar níveis de engajamento "
         "com maior objetividade do que o acompanhamento humano não instrumentado. No recorte desta PoC, a "
         "hipótese é operacionalizada de forma falseável: o modelo deve superar o baseline de classe majoritária "
         "em macro-F1 e kappa com intervalos de confiança de 95% não sobrepostos.", False),
    ])
    rico(doc, [
        ("Contribuições esperadas. ", True),
        ("(i) um protocolo de avaliação mais rigoroso que o predominante na literatura do DAiSEE, reportando "
         "macro-F1, kappa, revocação por classe e intervalos de confiança, e não apenas acurácia; (ii) evidência "
         "empírica sobre quanto do sinal de engajamento é recuperável a partir de marcos faciais leves, "
         "executáveis em hardware comum e com preservação de privacidade; (iii) uma ablação explícita de "
         "estratégias de balanceamento e de formulação ordinal, dimensionando o quanto o desbalanceamento – e não "
         "a arquitetura – limita o desempenho nesta tarefa.", False),
    ])

    # ---------------- 3. Estado da arte ----------------
    titulo_secao(doc, "3. Estado da Arte e Trabalhos Relacionados")
    par(doc,
        "O benchmark original do DAiSEE [1] estabeleceu 57,9% de acurácia com uma rede LRCN. A evolução seguinte "
        "combinou extratores espaciais profundos com modelagem temporal: o DFSTN [7] alcançou 58,8% com "
        "SE-ResNet-50 e LSTM com atenção, e o 3D DenseAttNet [8] chegou a 63,6% com autoatenção sobre cubos "
        "espaço-temporais. Duas linhas se mostraram mais produtivas que o simples aumento de capacidade. A "
        "primeira trata o rótulo como ordinal: Abedi e Khan [5] obtiveram 67,4% modelando engajamento como "
        "variável ordenada, e o trabalho contrastivo ordinal de Safa et al. [14] ataca diretamente a combinação de "
        "ordinalidade com classes raras. A segunda questiona a necessidade de modelar a ordem temporal: o Bag of "
        "States [9] atingiu 66,6% tratando o clipe como conjunto não ordenado de estados, resultado que enfraquece "
        "a premissa de que arquiteturas sequenciais pesadas sejam indispensáveis.")
    par(doc,
        "Para o recorte desta PoC, a evidência mais relevante é a de que representações compactas bastam. Abedi e "
        "Khan [6] usaram exclusivamente marcos do MediaPipe FaceMesh alimentando uma ST-GCN, obtendo desempenho de "
        "ponta em bases correlatas com custo baixo e sem trafegar imagens de rosto – propriedade de privacidade "
        "relevante em contexto educacional. Sugihdharma e Bachtiar [10] reportaram a maior acurácia revisada "
        "usando apenas marcos oculares e gaze, embora o valor (95,8%) destoe tanto da faixa dos demais trabalhos "
        "(58% a 73%) que seu protocolo precisa ser verificado antes de servir como cota superior. Em contrapartida, "
        "a fusão multimodal tem custo-benefício desfavorável: o VisioPhysioENet [11] adiciona sinais fisiológicos "
        "rPPG e fica em 63,1%. O estado da arte reportado é o ViBED-Net [13], com 73,4%, e o EngageFormer [12] "
        "atinge 63,9% – ambos com custo de treino muito acima do orçamento desta PoC. Por fim, um benchmark "
        "zero-shot de modelos de fundação [15] obteve kappa inferior a 0,10 no DAiSEE, com colapso das predições "
        "em uma única classe, o que justifica investir em treino supervisionado em vez de prompting.")
    rico(doc, [
        ("Diferenciação. ", True),
        ("Quase todos os trabalhos acima reportam acurácia como métrica central, o que é enganoso quando duas "
         "classes concentram mais de 94% dos exemplos: responder sempre “engajado” já rende cerca de 49% de "
         "acurácia. Poucos relatam kappa, macro-F1 ou revocação das classes raras, e menos ainda reportam "
         "intervalos de confiança. Este projeto não busca superar o número absoluto do ViBED-Net, e sim medir "
         "honestamente o que atributos faciais leves entregam sob um protocolo que torne visível o efeito do "
         "desbalanceamento – a lacuna metodológica que a área vem carregando.", False),
    ])

    # ---------------- 4. Objetivos ----------------
    titulo_secao(doc, "4. Objetivos")
    rico(doc, [
        ("Objetivo geral. ", True),
        ("Avaliar se indicadores visuais permitem identificar diferentes níveis de engajamento de estudantes em "
         "aulas síncronas de EaD. ", False),
        ("Critério de verificação: ", True),
        ("o melhor modelo deve superar o baseline de classe majoritária em macro-F1 e kappa, com intervalos de "
         "confiança de 95% (bootstrap, 1.000 reamostragens) não sobrepostos.", False),
    ])
    rico(doc, [
        ("Objetivo específico 1. ", True),
        ("Identificar os indicadores visuais mais relacionados ao engajamento, ao desengajamento e à distração. ",
         False),
        ("Critério: ", True),
        ("produzir um ranking de importância de atributos com estabilidade do top-5 em pelo menos 4 de 5 folds de "
         "validação cruzada estratificada.", False),
    ])
    rico(doc, [
        ("Objetivo específico 2. ", True),
        ("Avaliar a capacidade de modelos baseados em marcos faciais de diferenciar estudantes engajados de "
         "desengajados em vídeos reais. ", False),
        ("Critério: ", True),
        ("na formulação binária (Desengajado = níveis 0 e 1; Engajado = níveis 2 e 3), atingir revocação da classe "
         "minoritária maior ou igual a 60% e AUC-ROC maior ou igual a 0,70, contra revocação próxima de zero do "
         "baseline trivial.", False),
    ])
    par(doc,
        "Os três critérios são definidos a priori e são falseáveis: qualquer um dos desfechos – hipótese "
        "corroborada, parcialmente corroborada ou não corroborada – constitui resultado reportável, e a "
        "delimitação a uma PoC sobre dataset público mantém os objetivos alcançáveis no prazo do mestrado.")

    # ---------------- 5. Metodologia ----------------
    titulo_secao(doc, "5. Metodologia, Plano de Trabalho e Cronograma")
    rico(doc, [
        ("Dados e pré-processamento. ", True),
        (f"Utilizou-se o DAiSEE [1], com clipes de 10 s a 30 fps e resolução 640×480, rotulados de 0 a 3 em "
         f"Engajamento, Tédio, Confusão e Frustração. Foram respeitados os splits oficiais, disjuntos por sujeito "
         f"(69, 19 e 21 sujeitos em treino, validação e teste), condição necessária para comparabilidade com a "
         f"literatura. Dos 8.571 clipes rotulados nos três splits, {mil(meta['n_total'])} foram localizados em "
         f"disco e processados. De cada clipe amostraram-se 20 quadros uniformemente espaçados – e não todos os "
         f"quadros – e de cada quadro extraíram-se, via MediaPipe FaceMesh com refinamento de íris [17]: EAR "
         f"(abertura ocular), MAR (abertura bucal), pose da cabeça (yaw, pitch e roll por solvePnP) e posição "
         f"normalizada da íris (gaze horizontal, vertical e desvio em relação ao centro). Quadros sem rosto "
         f"detectado recebem marcação explícita, e a taxa de detecção entra como atributo; a detecção facial "
         f"média foi de {pct(meta['taxa_deteccao_media'])}. Os atributos por quadro são agregados por clipe "
         f"(média, desvio, mínimo e máximo), gerando {meta['n_features']} atributos tabulares. O ângulo de roll "
         f"exigiu correção de descontinuidade circular: por concentrar-se junto de ±180°, seu desvio-padrão vinha "
         f"inflado em 25,4% dos clipes.", False),
    ])
    rico(doc, [
        ("Modelagem. ", True),
        ("A Trilha A treina Random Forest, XGBoost e SVM sobre os atributos agregados, com GridSearchCV e "
         "validação cruzada estratificada dentro do treino, comparando quatro tratamentos de desbalanceamento: "
         "nenhum, class_weight balanceado, SMOTE e ponderação por inverso de frequência. A Trilha B treina uma "
         "BiLSTM rasa sobre a sequência de 20 passos por clipe, com perda focal ponderada por classe e parada "
         "antecipada monitorando macro-F1 de validação, além de uma variante ordinal CORAL [16] que substitui a "
         "saída softmax por limiares cumulativos. A extração de marcos é feita em CPU (a inferência do MediaPipe "
         "não usa GPU nesta configuração) e a Trilha B é treinada em GPU NVIDIA RTX 4060; ambos os modelos são "
         "suficientemente leves para execução em hardware de uso pessoal. A semente aleatória é fixada em 42 e as "
         "versões das bibliotecas são congeladas para garantir reprodutibilidade.", False),
    ])
    rico(doc, [
        ("Resultados. ", True),
        (f"A tabela abaixo consolida o desempenho no conjunto de teste ({mil(meta['n_test'])} clipes).", False),
    ])

    # Nomes de exibição acentuados (os CSVs dos notebooks usam ASCII)
    NOMES = {
        "Baseline trivial (classe majoritaria)": "Baseline trivial (classe majoritária)",
        "RandomForest (sem balanceamento)": "Random Forest (sem balanceamento)",
        "RandomForest (class_weight=balanced)": "Random Forest (class_weight=balanced)",
        "RandomForest + SMOTE": "Random Forest + SMOTE",
        "XGBoost (sample_weight balanceado)": "XGBoost (ponderação por classe)",
        "SVM (class_weight=balanced)": "SVM (class_weight=balanced)",
        "BiLSTM + Focal Loss (class-weighted)": "BiLSTM + perda focal (ponderada)",
        "BiLSTM ordinal (CORAL)": "BiLSTM ordinal (CORAL)",
    }
    linhas = []
    for _, r in cons.iterrows():
        nome = str(r["modelo"]).strip()
        linhas.append([
            NOMES.get(nome, nome),
            n(r["acc"]), n(r["macro_f1"]), n(r["kappa"]),
        ])
    tabela(doc, ["Modelo", "Acurácia", "Macro-F1", "Kappa"], linhas, larguras=[9.0, 2.4, 2.4, 2.4])

    par(doc,
        f"O bootstrap com 1.000 reamostragens sobre o teste situa o melhor modelo em macro-F1 "
        f"{n(b_model['macro_f1_media'])} (IC95% [{n(b_model['macro_f1_ic_low'])}; {n(b_model['macro_f1_ic_high'])}]) "
        f"e kappa {n(b_model['kappa_media'])} (IC95% [{n(b_model['kappa_ic_low'])}; {n(b_model['kappa_ic_high'])}]), "
        f"contra macro-F1 {n(b_base['macro_f1_media'])} (IC95% [{n(b_base['macro_f1_ic_low'])}; "
        f"{n(b_base['macro_f1_ic_high'])}]) e kappa {n(b_base['kappa_media'])} do baseline trivial. "
        f"{meta['frase_ic']} No Objetivo 2, a formulação binária alcançou revocação de "
        f"{pct(binr['recall_desengajado'], 1)} para a classe Desengajado e AUC-ROC de {n(binr['auc_roc'])}. "
        f"Quanto ao Objetivo 1, o top-5 de atributos foi {', '.join(top5)}, dos quais "
        f"{len(estaveis)} permaneceram no top-5 em pelo menos 4 dos 5 folds. Vale registrar que a formulação "
        f"ordinal CORAL superou a variante nominal com perda focal na Trilha B (macro-F1 de 0,266 contra 0,121), "
        f"indicando que preservar a ordem dos níveis ajuda mais que reponderar classes de forma agressiva.",
        space_after=4)
    par(doc, meta["paragrafo_discussao"], space_after=4)
    rico(doc, [
        ("Limitações. ", True),
        (f"(i) a amostragem de 20 quadros por clipe descarta microdinâmicas rápidas; (ii) o conjunto de teste "
         f"contém apenas {meta['n_classe0_teste']} clipes da classe 0, o que torna o macro-F1 sensível a poucos "
         f"acertos e amplia os intervalos de confiança; (iii) marcos faciais descartam contexto de cena e postura "
         f"corporal, usados pelos modelos de ponta [13]; (iv) a PoC restringiu-se deliberadamente a atributos "
         f"geométricos leves, de modo que backbones visuais pré-treinados – viáveis no hardware disponível – "
         f"ainda não foram explorados; (v) há indícios de ruído de rótulo por sujeito, já documentados na "
         f"literatura.", False),
    ])
    rico(doc, [
        ("Plano de trabalho e cronograma. ", True),
        ("A PoC foi organizada em seis etapas reprodutíveis, já executadas: (1) preparação de ambiente e validação "
         "do dataset; (2) análise exploratória; (3) extração de marcos faciais; (4) linha de base clássica e "
         "ranking de atributos; (5) modelo temporal e variante ordinal; (6) avaliação comparativa e bootstrap. No "
         "horizonte do mestrado, seguem-se: meses 1 a 3, validação cruzada em bases adicionais (EngageNet, Student "
         "Engagement Dataset) para medir generalização entre domínios; meses 4 a 7, contexto de cena via backbone "
         "pré-treinado e aprendizado contrastivo ordinal [14], já viáveis na GPU local; meses 8 a 10, ruído de "
         "rótulo por sujeito e calibração de incerteza; meses 11 a 14, protótipo de painel de feedback ao docente "
         "e avaliação com professores; meses 15 a 18, redação e defesa.", False),
    ])

    # ---------------- 6. Referencias ----------------
    titulo_secao(doc, "6. Referências")
    for i, r in enumerate(REFERENCIAS, start=1):
        par(doc, f"[{i}] {r}", align="left", space_after=1)

    return doc


def main():
    d = carregar_insumos()

    cons = d["consolidado"]
    dec = d["decisao"].iloc[0]
    boot = d["boot"]
    binr = d["binario"].iloc[0]

    FEAT = ROOT / "datasets" / "DAiSEE" / "features"
    taxa = pd.read_csv(FEAT / "taxa_deteccao_rosto.csv").iloc[0]
    nclip = pd.read_csv(FEAT / "n_clipes_por_split.csv").iloc[0]
    pred_a = pd.read_csv(FEAT / "predicoes_teste_trilha_a.csv")

    b_model, b_base = boot.iloc[0], boot.iloc[1]
    sobrepoe = (b_model["macro_f1_ic_low"] <= b_base["macro_f1_ic_high"]) or \
               (b_model["kappa_ic_low"] <= b_base["kappa_ic_high"])

    frase_ic = ("Os intervalos não se sobrepõem, caracterizando ganho estatisticamente significativo sobre o "
                "baseline." if not sobrepoe else
                "Os intervalos se sobrepõem, de modo que o ganho sobre o baseline não é estatisticamente "
                "conclusivo.")

    veredito = dec["veredito_hipotese"]
    if veredito == "CORROBORADA":
        disc = ("O conjunto de evidências sustenta a hipótese: marcos faciais leves, sem transmissão de imagem, "
                "recuperam sinal de engajamento muito acima do acaso e identificam a maior parte dos estudantes "
                "desengajados, que são justamente o alvo pedagógico da intervenção.")
    elif veredito == "PARCIALMENTE CORROBORADA":
        disc = ("O desfecho é de corroboração parcial: há ganho estatisticamente significativo sobre o baseline em "
                "macro-F1 e kappa, confirmando que os marcos faciais carregam sinal real de engajamento, mas a "
                "revocação da classe Desengajado ficou muito abaixo da meta de 60%. O fator limitante, portanto, "
                "não é a ausência de sinal, e sim o desbalanceamento extremo: mesmo com ponderação de classes, "
                "SMOTE e perda focal, as classes raras permanecem sub-representadas a ponto de nenhuma das "
                "mitigações testadas as recuperar plenamente. Um kappa de 0,11 indica concordância apenas "
                "ligeiramente superior ao acaso, insuficiente para uso pedagógico direto, ainda que suficiente "
                "para demonstrar que o sinal existe. Esse é precisamente o diagnóstico que a acurácia, métrica "
                "dominante na literatura, tornaria invisível: o baseline trivial atinge 51,8% de acurácia, valor "
                "que isolado pareceria competitivo.")
    else:
        disc = ("O desfecho não corrobora a hipótese no recorte testado: o modelo não superou o baseline trivial "
                "com significância estatística. A leitura mais provável é que atributos geométricos agregados por "
                "clipe sejam insuficientes para o rótulo de engajamento do DAiSEE, seja por perda de informação na "
                "agregação, seja por ruído de anotação – hipóteses que o próximo ciclo deverá separar.")

    y_true = pred_a["y_true"]
    meta = {
        "n_total": int(nclip["train"] + nclip["validation"] + nclip["test"]),
        "n_test": int(nclip["test"]),
        "n_features": len(pd.read_csv(FEAT / "importancias_features_engagement.csv")),
        "taxa_deteccao_media": float((taxa["train"] * nclip["train"] + taxa["validation"] * nclip["validation"] +
                                      taxa["test"] * nclip["test"]) / (nclip["train"] + nclip["validation"] + nclip["test"])),
        "n_classe0_teste": int((y_true == 0).sum()),
        "frase_ic": frase_ic,
        "paragrafo_discussao": disc,
    }
    d["meta"] = meta

    doc = montar_documento(d)
    doc.save(str(SAIDA))
    print("Documento gerado:", SAIDA)
    return SAIDA


if __name__ == "__main__":
    main()
